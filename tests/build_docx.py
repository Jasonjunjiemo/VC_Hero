"""合成《项目原型架构性能测试.docx》到项目根目录。"""
import json
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DIAG = os.path.join(ROOT, "docs", "diagrams")
PERF = json.load(open(os.path.join(ROOT, "tests", "_perf_result.json"), encoding="utf-8"))

doc = Document()
# 文档元数据（避免默认 python-docx 作者名）
cp = doc.core_properties
cp.author = "Jasonjunjiemo"
cp.last_modified_by = "Jasonjunjiemo"
cp.title = "VC Hero 项目原型架构性能测试"

# 全局中文字体
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def set_cn(run, size=None, bold=None, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def h(text, level=1):
    p = doc.add_heading("", level=level)
    r = p.add_run(text)
    set_cn(r, size={0: 20, 1: 15, 2: 12.5}.get(level, 12), bold=True,
           color=(0x1A, 0x22, 0x33) if level else (0x46, 0x56, 0xE0))
    return p


def para(text, size=10.5, bold=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_cn(r, size=size, bold=bold, color=color)
    return p


def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_cn(r, size=9, color=(0x8A, 0x94, 0xA6))


def img(path, width_cm=15.5, cap=None):
    doc.add_picture(path, width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cap:
        caption(cap)


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, htxt in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(htxt)
        set_cn(r, size=9.5, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(val))
            set_cn(r, size=9.5)
    doc.add_paragraph()
    return t


# ================= 封面标题 =================
h("VC Hero · 项目原型架构性能测试", level=0)
para("AI 虚拟面试官（早期 VC Deal Sourcing 面试陪练） —— 原型图 / 架构图 / 性能测试三部分汇总。"
     "线上地址：http://152.136.60.146 · 仓库：https://github.com/Jasonjunjiemo/VC_Hero", size=9.5,
     color=(0x4C, 0x56, 0x70))
doc.add_paragraph()

# ================= 一、原型图 =================
h("一、原型图", level=1)
para("以下线框图为当前线上版本的真实布局抽象：顶栏为透明浮动控件（无横条），双 Tab 主页、"
     "对话式面试页、可展开/收起的结果面板。")

h("1.1 主页（AI 面试官 / AI 训练官）", level=2)
img(os.path.join(DIAG, "proto1_home.png"),
    cap="图 1-1 主页：浮动 logo 与用户控件、双 Tab 切换器、hero 介绍区、会话管理面板")

h("1.2 面试对话页", level=2)
img(os.path.join(DIAG, "proto2_chat.png"),
    cap="图 1-2 对话页：浮动控件（返回/标题/计时/结束）、消息流、一体化输入区、右侧结果速览抽屉")

h("1.3 面试结果面板", level=2)
img(os.path.join(DIAG, "proto3_result.png"),
    cap="图 1-3 结果面板：展开态（评级圆环+五维分数条+反馈正文）与收起态（仅头部一行，点击展开）")

# ================= 二、架构图 =================
h("二、架构图", level=2 - 1)
para("技术栈：Flask + 原生 HTML/CSS/JS（单页应用 + Jinja 服务端渲染首屏）、本地文件系统存储、"
     "Kimi 开放平台 API（kimi-k3，关闭推理）、全局限速与 FIFO 负载均衡。全部提示词以 Markdown "
     "存于 prompts/，后端零硬编码。")
img(os.path.join(DIAG, "architecture.png"),
    cap="图 2-1 系统架构：浏览器 ↔ Flask（双引擎） ↔ 文件存储 / Kimi API；限速器统一调度全部 LLM 调用")

h("2.1 后端模块", level=2)
table(["模块", "职责"], [
    ["routes.py", "REST API；/start /answer /finish /continue 按会话 kind 分发到面试或训练引擎"],
    ["interview.py", "面试状态机：简历问答（模型以 [[CV_DONE]] 自行收尾）→ 场景任务 → 评分（A-E）"],
    ["training.py", "训练引擎：自由回合场景训练，结束生成训练总结；支持链式继续"],
    ["kimi.py", "Kimi 客户端：生成 → 恰好 2 轮自我迭代 → 才发出；429/空内容退避重试"],
    ["ratelimit.py", "滑动窗口限速（token/h、response/h、¥5/h）+ FIFO 公平队列负载均衡，预占-实结核算"],
    ["storage.py / pdfutil.py", "文件系统 JSON 存储（原子写入）；PDF/TXT/MD/DOCX 文本抽取"],
    ["prompts/（7 个 md）", "三份原始规则 + 总控 + 自检×2 + 评分格式 + 任务场景库"],
])

h("2.2 关键交互流", level=2)
para("① 打开活动会话页 → 服务端渲染首屏（顶栏/消息/输入框一次到位），前端水合接管交互；"
     "② 发送回答 → 乐观回显 → 三点思考气泡 → 考官回复（每轮 = 生成 + 恰好 2 轮自检，共 3 次 API 调用）；"
     "③ 结束/聊满 6 轮 → 生成结果面板（可收起、可经右侧速览栏跳转）；④ 点击继续 → 链式续聊，"
     "可多次产生结果快照。")

# ================= 三、性能测试 =================
h("三、性能测试", level=1)
para("测试对象：本地实例（Windows 11，Python 3.12，Flask 开发服务器 threaded 模式），"
     "测试脚本 tests/perf_test.py 实测（串行 100 次取 p50/p95/max；并发 30 线程 × 20 次；"
     "AI 一轮 = 生成 + 2 轮自检共 3 次 Kimi 调用）。", size=9.5)

h("3.1 静态与轻量 API 延迟（毫秒）", level=2)
rows = []
for k in ["A1 GET / 首页(SSR模板)", "A2 GET /static/app.js", "A3 GET /api/health",
          "B1 GET /api/me", "B2 GET /api/sessions?kind=interview"]:
    r = PERF[k]
    rows.append([k.replace("GET ", "GET "), r["n"], r["p50"], r["p95"], r["mean"], r["max"]])
table(["端点", "样本数", "p50", "p95", "均值", "最大"], rows)

h("3.2 服务端渲染会话首屏（毫秒）", level=2)
r = PERF["D1 GET /session/<id> (SSR首屏)"]
table(["端点", "样本数", "p50", "p95", "均值", "最大"],
      [["GET /session/<id>（含会话数据读取+模板渲染）", r["n"], r["p50"], r["p95"], r["mean"], r["max"]]])

h("3.3 并发能力", level=2)
rows = []
for k in ["C1 并发 GET /api/health", "C2 并发 GET / (SSR)"]:
    r = PERF[k]
    rows.append([k, f"{r['threads']} 线程×20 次 = {r['requests']} 请求",
                 r["wall_s"], r["req_per_s"], r["p50"], r["p95"]])
table(["场景", "规模", "总耗时(s)", "吞吐(req/s)", "p50(ms)", "p95(ms)"], rows)
para("注：并发测试的瓶颈在 Python 测试客户端（urllib + GIL 与连接开销）；服务端 600 个请求"
     "全部成功返回、无错误。生产部署（Ubuntu + 同版本 Flask）行为一致，生产级部署可替换为 "
     "waitress/gunicorn 以获得更高并发吞吐。", size=9, color=(0x4C, 0x56, 0x70))

h("3.4 AI 一轮回复延迟（秒）", level=2)
r = PERF["E1 AI一轮回复(生成+2轮自检)"]
table(["指标", "值"],
      [["样本数", r["n"]],
       ["均值", f'{r["mean_s"]} s'],
       ["最小", f'{r["min_s"]} s'],
       ["最大", f'{r["max_s"]} s'],
       ["说明", "kimi-k3 关闭推理（thinking disabled, temperature 0.6）；"
                "单轮 API 5-12s，一轮回复 3 次调用合计约 20-25s"]])

h("3.5 结论", level=2)
para("① 非 AI 路径全部在 50ms 内（p50 约 15-20ms），SSR 首屏 p50 45ms，体验流畅；"
     "② AI 交互延迟由 Kimi API 决定（每轮约 20-25 秒，3 次调用），前端以思考气泡+乐观回显覆盖等待；"
     "③ 全局限速（1 小时窗口：40 万 token / 200 次响应 / ¥5）在压测期间未触发排队，多人场景下"
     "超限请求按 FIFO 公平队列等待而非报错。")

doc.save(os.path.join(ROOT, "项目原型架构性能测试.docx"))
print("saved:", os.path.join(ROOT, "项目原型架构性能测试.docx"))
